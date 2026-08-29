from __future__ import annotations

import argparse
import re
import sys
import zipfile
from pathlib import Path
from typing import Iterable, List, Sequence, Tuple

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "prompt_docs_sources"
OUTPUT_DIR = ROOT / "deliverables"

FILES = [
    (
        SOURCE_DIR / "prompt_with_ruby.md",
        OUTPUT_DIR / "高考日语新题型配套解析最终Prompt（含标音版）.docx",
        "高考日语新题型配套解析最终Prompt（含标音版）",
    ),
    (
        SOURCE_DIR / "prompt_without_ruby.md",
        OUTPUT_DIR / "高考日语新题型配套解析最终Prompt（不含标音版）.docx",
        "高考日语新题型配套解析最终Prompt（不含标音版）",
    ),
]

CN_FONT = "宋体"
JP_FONT = "MS Mincho"
BODY_SIZE = 11
RUBY_SIZE = 5.5

JP_KANA_RE = re.compile(r"[ぁ-ゟァ-ヿー]|")
JP_QUOTED_RE = re.compile(r"(「[^」]*」|『[^』]*』|`[^`]*`)")
BOLD_RE = re.compile(r"(\*\*[^*]+\*\*)")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, font_name: str, size: float = BODY_SIZE, bold: bool | None = None,
                 italic: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font_name)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "ja-JP" if font_name == JP_FONT else "zh-CN")
    lang.set(qn("w:eastAsia"), "ja-JP" if font_name == JP_FONT else "zh-CN")
    no_proof = r_pr.find(qn("w:noProof"))
    if no_proof is None:
        r_pr.append(OxmlElement("w:noProof"))


def looks_japanese(text: str, force_quoted: bool = False) -> bool:
    if force_quoted:
        return True
    return bool(re.search(r"[ぁ-ゟァ-ヿー]", text))


def split_japanese(text: str) -> List[Tuple[str, bool]]:
    """Split text so Japanese quoted/code spans and kana-containing fragments use MS Mincho."""
    if not text:
        return []
    parts: List[Tuple[str, bool]] = []
    pos = 0
    for m in JP_QUOTED_RE.finditer(text):
        if m.start() > pos:
            plain = text[pos:m.start()]
            parts.extend(split_kana_fragments(plain))
        quoted = m.group(0)
        parts.append((quoted, True))
        pos = m.end()
    if pos < len(text):
        parts.extend(split_kana_fragments(text[pos:]))
    return [(t, jp) for t, jp in parts if t]


def split_kana_fragments(text: str) -> List[Tuple[str, bool]]:
    # Treat continuous tokens containing kana as Japanese, while leaving Chinese prose in SimSun.
    tokens = re.split(r"(\s+|[，。；：！？、（）【】《》“”‘’])", text)
    out: List[Tuple[str, bool]] = []
    for token in tokens:
        if not token:
            continue
        out.append((token, looks_japanese(token)))
    return out


def add_inline(paragraph, text: str, default_size: float = BODY_SIZE,
               default_bold: bool = False, color: RGBColor | None = None) -> None:
    chunks = BOLD_RE.split(text)
    for chunk in chunks:
        if not chunk:
            continue
        is_bold = chunk.startswith("**") and chunk.endswith("**")
        clean = chunk[2:-2] if is_bold else chunk
        for segment, is_jp in split_japanese(clean):
            run = paragraph.add_run(segment)
            set_run_font(
                run,
                JP_FONT if is_jp else CN_FONT,
                default_size,
                bold=(default_bold or is_bold),
                color=color,
            )


def set_paragraph_format(paragraph, line_spacing: float = 1.28, after: float = 3,
                         before: float = 0, keep_with_next: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing
    fmt.space_after = Pt(after)
    fmt.space_before = Pt(before)
    fmt.keep_with_next = keep_with_next
    fmt.widow_control = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, CN_FONT, 9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()._r
    r.append(fld_char1)
    r.append(instr_text)
    r.append(fld_char2)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, CN_FONT, 9)


def configure_document(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.15)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = CN_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    normal.paragraph_format.line_spacing = 1.28
    normal.paragraph_format.space_after = Pt(3)

    for style_name, size, color in [
        ("Title", 22, "17365D"),
        ("Heading 1", 16, "17365D"),
        ("Heading 2", 14, "1F4E79"),
        ("Heading 3", 12, "2F5597"),
        ("Heading 4", 11, "365F91"),
    ]:
        style = styles[style_name]
        style.font.name = CN_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(title)
    set_run_font(run, CN_FONT, 9, color=RGBColor(110, 110, 110))
    add_page_number(section.footer.paragraphs[0])

    doc.core_properties.title = title
    doc.core_properties.subject = "高考日语新题型配套答案及解析出版级制作标准"
    doc.core_properties.author = "阿狸老师日语课堂"
    doc.core_properties.keywords = "高考日语, 单选, 阅读, 语篇, 完形, 解析, Prompt"


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style="Title" if level == 0 else f"Heading {min(level, 4)}")
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(30)
        p.paragraph_format.space_after = Pt(16)
    add_inline(p, text, default_size=22 if level == 0 else {1: 16, 2: 14, 3: 12, 4: 11}.get(level, 11), default_bold=True)
    set_paragraph_format(p, line_spacing=1.15, after=8 if level == 0 else 5, keep_with_next=True)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B4C6E7")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    set_paragraph_format(p, after=5)


def add_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if i == 0:
                set_cell_shading(cell, "D9EAF7")
            p = cell.paragraphs[0]
            add_inline(p, value, default_size=10.5, default_bold=(i == 0))
            set_paragraph_format(p, line_spacing=1.12, after=0)
    doc.add_paragraph()


def parse_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    first_title_seen = False
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped:
            p = doc.add_paragraph()
            set_paragraph_format(p, after=1)
            i += 1
            continue

        if stripped == "---":
            add_horizontal_rule(doc)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows: List[List[str]] = []
            for idx, line in enumerate(table_lines):
                cells = [c.strip() for c in line.strip("|").split("|")]
                if idx == 1 and all(re.fullmatch(r":?-{2,}:?", c or "-") for c in cells):
                    continue
                rows.append(cells)
            add_table(doc, rows)
            continue

        if stripped.startswith("# "):
            add_heading(doc, stripped[2:].strip(), 0 if not first_title_seen else 1)
            first_title_seen = True
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 1)
            i += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 2)
            i += 1
            continue
        if stripped.startswith("#### "):
            add_heading(doc, stripped[5:].strip(), 3)
            i += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.right_indent = Cm(0.4)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F2F6FA")
            p_pr.append(shd)
            add_inline(p, stripped[2:].strip(), default_size=10.5)
            set_paragraph_format(p, line_spacing=1.22, after=4, before=2)
            i += 1
            continue

        if re.match(r"^- \[[ xX]\] ", stripped):
            checked = stripped[3].lower() == "x"
            text = stripped[6:]
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, ("☑ " if checked else "☐ ") + text)
            set_paragraph_format(p, after=2)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:].strip())
            set_paragraph_format(p, after=2)
            i += 1
            continue

        if re.match(r"^\d+[\.、]", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.15)
            add_inline(p, stripped)
            set_paragraph_format(p, after=2)
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, stripped)
        set_paragraph_format(p, after=3)
        i += 1


def build_document(source: Path, output: Path, title: str) -> None:
    markdown = source.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc, title)
    parse_markdown(doc, markdown)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def iter_doc_text(doc: Document) -> Iterable[str]:
    for p in doc.paragraphs:
        yield p.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p.text


def validate_docx(path: Path, expected_title: str) -> List[str]:
    errors: List[str] = []
    if not path.exists() or path.stat().st_size < 20000:
        errors.append(f"File missing or unexpectedly small: {path}")
        return errors
    if not zipfile.is_zipfile(path):
        errors.append(f"Not a valid DOCX ZIP package: {path}")
        return errors
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
        if bad:
            errors.append(f"Corrupt ZIP member {bad} in {path}")
        required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
        missing = required - set(zf.namelist())
        if missing:
            errors.append(f"Missing DOCX members {sorted(missing)} in {path}")
    doc = Document(path)
    text = "\n".join(iter_doc_text(doc))
    if expected_title not in text:
        errors.append(f"Title not found in {path}")
    if "�" in text or "\uFFFD" in text:
        errors.append(f"Replacement character found in {path}")
    if "TOC \\o" in text:
        errors.append(f"Raw TOC field code found in visible text: {path}")
    if len(text) < 15000:
        errors.append(f"Document text unexpectedly short ({len(text)} chars): {path}")
    return errors


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--validate-only", action="store_true")
    args = parser.parse_args()

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if not args.validate_only:
        for source, output, title in FILES:
            build_document(source, output, title)

    all_errors: List[str] = []
    report_lines = ["高考日语新题型配套解析 Prompt｜自动结构校验报告", ""]
    for _, output, title in FILES:
        errors = validate_docx(output, title)
        if errors:
            all_errors.extend(errors)
            report_lines.append(f"[FAIL] {output.name}")
            report_lines.extend(f"  - {e}" for e in errors)
        else:
            report_lines.append(f"[PASS] {output.name}")
            report_lines.append(f"  - size: {output.stat().st_size} bytes")
    report = OUTPUT_DIR / "validation_report.txt"
    report.write_text("\n".join(report_lines) + "\n", encoding="utf-8")
    print(report.read_text(encoding="utf-8"))
    return 1 if all_errors else 0


if __name__ == "__main__":
    raise SystemExit(main())
