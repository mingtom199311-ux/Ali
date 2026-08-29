from __future__ import annotations

import argparse
import re
import zipfile
from pathlib import Path
from typing import Iterable, List, Sequence, Tuple

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "prompt_docs_sources"
OUTPUT_DIR = ROOT / "deliverables_v2"

FILES = [
    (
        SOURCE_DIR / "prompt_with_ruby.md",
        OUTPUT_DIR / "高考日语新题型配套解析最终Prompt（含标音版）.docx",
        "高考日语新题型配套解析最终 Prompt（含标音版）",
    ),
    (
        SOURCE_DIR / "prompt_without_ruby.md",
        OUTPUT_DIR / "高考日语新题型配套解析最终Prompt（不含标音版）.docx",
        "高考日语新题型配套解析最终 Prompt（不含标音版）",
    ),
]

CN_FONT = "宋体"
JP_FONT = "MS Mincho"
BODY_SIZE = 11

BOLD_RE = re.compile(r"(\*\*[^*]+\*\*)")
JP_QUOTE_RE = re.compile(r"(「[^」]*」|『[^』]*』|`[^`]*`)")
KANA_RE = re.compile(r"[ぁ-ゟァ-ヿー]")


def set_run_font(
    run,
    font_name: str,
    size: float = BODY_SIZE,
    *,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font_name)

    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang_code = "ja-JP" if font_name == JP_FONT else "zh-CN"
    lang.set(qn("w:val"), lang_code)
    lang.set(qn("w:eastAsia"), lang_code)

    if r_pr.find(qn("w:noProof")) is None:
        r_pr.append(OxmlElement("w:noProof"))


def split_kana_tokens(text: str) -> List[Tuple[str, bool]]:
    if not text:
        return []
    # Keep punctuation/whitespace as independent segments; tokens containing kana are Japanese.
    tokens = re.split(r"(\s+|[，。；：！？、（）【】《》“”‘’])", text)
    return [(token, bool(KANA_RE.search(token))) for token in tokens if token]


def split_language_segments(text: str) -> List[Tuple[str, bool]]:
    """Return (segment, is_japanese) pairs.

    Chinese prose remains SimSun. Japanese is recognized conservatively by kana.
    Quoted/code spans are Japanese only when they contain kana or a Japanese grammar marker.
    This avoids treating Chinese Han characters as Japanese merely because they are CJK.
    """
    if not text:
        return []
    result: List[Tuple[str, bool]] = []
    pos = 0
    for match in JP_QUOTE_RE.finditer(text):
        if match.start() > pos:
            result.extend(split_kana_tokens(text[pos:match.start()]))
        quoted = match.group(0)
        is_japanese = bool(KANA_RE.search(quoted)) or bool(
            re.search(r"\b(?:V|N|A|AN)-|\+|ます|て|た|ば|そう|よう|らしい", quoted)
        )
        result.append((quoted, is_japanese))
        pos = match.end()
    if pos < len(text):
        result.extend(split_kana_tokens(text[pos:]))
    return [(segment, is_japanese) for segment, is_japanese in result if segment]


def add_inline(
    paragraph,
    text: str,
    *,
    size: float = BODY_SIZE,
    bold: bool = False,
    color: RGBColor | None = None,
) -> None:
    for chunk in BOLD_RE.split(text):
        if not chunk:
            continue
        chunk_bold = chunk.startswith("**") and chunk.endswith("**")
        clean = chunk[2:-2] if chunk_bold else chunk
        for segment, is_japanese in split_language_segments(clean):
            run = paragraph.add_run(segment)
            set_run_font(
                run,
                JP_FONT if is_japanese else CN_FONT,
                size,
                bold=(bold or chunk_bold),
                color=color,
            )


def set_paragraph_format(
    paragraph,
    *,
    line_spacing: float = 1.26,
    before: float = 0,
    after: float = 3,
    keep_with_next: bool = False,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.keep_with_next = keep_with_next
    fmt.widow_control = True


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, CN_FONT, 9, color=RGBColor(115, 115, 115))

    field_run = paragraph.add_run()
    field_r = field_run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_r.extend([begin, instr, end])

    run = paragraph.add_run(" 页")
    set_run_font(run, CN_FONT, 9, color=RGBColor(115, 115, 115))


def configure_document(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.15)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = CN_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    normal.paragraph_format.line_spacing = 1.26
    normal.paragraph_format.space_after = Pt(3)

    style_specs = {
        "Title": (21, "17365D", 0, 14),
        "Heading 1": (16, "17365D", 10, 6),
        "Heading 2": (14, "1F4E79", 9, 5),
        "Heading 3": (12, "2F5597", 7, 4),
        "Heading 4": (11, "365F91", 6, 3),
    }
    for style_name, (size, color, before, after) in style_specs.items():
        style = doc.styles[style_name]
        style.font.name = CN_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run(title)
    set_run_font(header_run, CN_FONT, 9, color=RGBColor(115, 115, 115))
    add_page_number(section.footer.paragraphs[0])

    doc.core_properties.title = title
    doc.core_properties.subject = "高考日语新题型配套答案及解析出版级制作标准"
    doc.core_properties.author = "阿狸老师日语课堂"
    doc.core_properties.keywords = "高考日语, 单选, 阅读, 语篇, 完形, 解析, Prompt"


def add_heading(doc: Document, text: str, level: int) -> None:
    if level == 0:
        paragraph = doc.add_paragraph(style="Title")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(26)
        paragraph.paragraph_format.space_after = Pt(15)
        size = 21
    else:
        paragraph = doc.add_paragraph(style=f"Heading {min(level, 4)}")
        size = {1: 16, 2: 14, 3: 12, 4: 11}.get(level, 11)
    add_inline(paragraph, text, size=size, bold=True)
    set_paragraph_format(paragraph, line_spacing=1.12, after=8 if level == 0 else 4, keep_with_next=True)


def add_horizontal_rule(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B4C6E7")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    set_paragraph_format(paragraph, after=4)


def add_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            value = row[column_index] if column_index < len(row) else ""
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
            paragraph = cell.paragraphs[0]
            add_inline(paragraph, value, size=10.3, bold=(row_index == 0))
            set_paragraph_format(paragraph, line_spacing=1.12, after=0)
    doc.add_paragraph()


def normalize_source(markdown: str) -> str:
    # Avoid displaying an actual replacement glyph or a raw Word TOC command in the reusable prompt document.
    markdown = markdown.replace("Unicode替代字符「�」", "Unicode替代字符（U+FFFD）")
    markdown = markdown.replace("`TOC \\o...`", "`TOC目录域代码`")
    markdown = markdown.replace("`TOC \\o...`等域代码", "`TOC目录域代码`")
    return markdown.replace("\r\n", "\n").replace("\r", "\n")


def parse_markdown(doc: Document, markdown: str) -> None:
    lines = normalize_source(markdown).split("\n")
    first_title_seen = False
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        if stripped == "---":
            add_horizontal_rule(doc)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines: List[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows: List[List[str]] = []
            for line_index, line in enumerate(table_lines):
                cells = [cell.strip() for cell in line.strip("|").split("|")]
                if line_index == 1 and all(re.fullmatch(r":?-{2,}:?", cell or "-") for cell in cells):
                    continue
                rows.append(cells)
            add_table(doc, rows)
            continue

        if stripped.startswith("# "):
            add_heading(doc, stripped[2:].strip(), 0 if not first_title_seen else 1)
            first_title_seen = True
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 1)
            i += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 2)
            i += 1
            continue
        if stripped.startswith("#### "):
            add_heading(doc, stripped[5:].strip(), 3)
            i += 1
            continue

        if stripped.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.65)
            paragraph.paragraph_format.right_indent = Cm(0.35)
            p_pr = paragraph._p.get_or_add_pPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F2F6FA")
            p_pr.append(shading)
            add_inline(paragraph, stripped[2:].strip(), size=10.4)
            set_paragraph_format(paragraph, line_spacing=1.2, before=2, after=4)
            i += 1
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:].strip())
            set_paragraph_format(paragraph, after=2)
            i += 1
            continue

        if re.match(r"^\d+[\.、]", stripped):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.12)
            add_inline(paragraph, stripped)
            set_paragraph_format(paragraph, after=2)
            i += 1
            continue

        paragraph = doc.add_paragraph()
        add_inline(paragraph, stripped)
        set_paragraph_format(paragraph, after=3)
        i += 1


def build_document(source: Path, output: Path, title: str) -> None:
    doc = Document()
    configure_document(doc, title)
    parse_markdown(doc, source.read_text(encoding="utf-8"))
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def iter_visible_text(doc: Document) -> Iterable[str]:
    for paragraph in doc.paragraphs:
        yield paragraph.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph.text


def validate_docx(path: Path, expected_title: str) -> List[str]:
    errors: List[str] = []
    if not path.exists() or path.stat().st_size < 25000:
        return [f"文件缺失或体积异常：{path}"]
    if not zipfile.is_zipfile(path):
        return [f"不是有效DOCX压缩包：{path}"]

    with zipfile.ZipFile(path) as archive:
        bad_member = archive.testzip()
        if bad_member:
            errors.append(f"DOCX成员损坏：{bad_member}")
        required_members = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
        missing_members = required_members - set(archive.namelist())
        if missing_members:
            errors.append(f"缺少DOCX关键成员：{sorted(missing_members)}")
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="strict")
        styles_xml = archive.read("word/styles.xml").decode("utf-8", errors="strict")
        if CN_FONT not in document_xml + styles_xml:
            errors.append("未检测到宋体设置")
        if JP_FONT not in document_xml + styles_xml:
            errors.append("未检测到MS Mincho设置")
        if "\x00" in document_xml:
            errors.append("检测到NUL控制字符")

    doc = Document(path)
    visible_text = "\n".join(iter_visible_text(doc))
    if expected_title not in visible_text:
        errors.append("正文中未找到完整标题")
    if len(visible_text) < 7000:
        errors.append(f"正文长度异常：{len(visible_text)}字符")
    if visible_text.count("##") > 0 or visible_text.count("**") > 0:
        errors.append("检测到未解析的Markdown标记")
    if "U+FFFD" not in visible_text:
        errors.append("字符清洗规则中的U+FFFD说明缺失")
    if "开始执行指令" not in visible_text or "终审清单" not in visible_text:
        errors.append("关键章节缺失")
    return errors


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--validate-only", action="store_true")
    args = parser.parse_args()

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if not args.validate_only:
        for source, output, title in FILES:
            build_document(source, output, title)

    report_lines = ["高考日语新题型配套解析最终Prompt｜自动结构校验报告", ""]
    failed = False
    for _, output, title in FILES:
        errors = validate_docx(output, title)
        if errors:
            failed = True
            report_lines.append(f"[FAIL] {output.name}")
            report_lines.extend(f"  - {error}" for error in errors)
        else:
            report_lines.append(f"[PASS] {output.name}")
            report_lines.append(f"  - 文件体积：{output.stat().st_size} bytes")
    report = OUTPUT_DIR / "validation_report.txt"
    report.write_text("\n".join(report_lines) + "\n", encoding="utf-8")
    print(report.read_text(encoding="utf-8"))
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
