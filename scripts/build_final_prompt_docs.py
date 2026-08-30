from __future__ import annotations

from pathlib import Path
import re
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT_DIR = Path("dist")
OUT_DIR.mkdir(parents=True, exist_ok=True)

TITLE_RUBY = "高考日语新题型配套解析最终Prompt（含标音版）"
TITLE_PLAIN = "高考日语新题型配套解析最终Prompt（不含标音版）"

COMMON_PROMPT = r'''
# 高考日语新题型配套解析｜出版级制作与审校最终 Prompt

你现在是一名“高考日语新题型出版级答案解析制作与审校专家”。请根据我提供的【题目文档】【旧答案及解析文档】【官方答案或题源】【原始讲义／知识资料】【版式母本】完成配套答案与解析的制作、审核或重做。

本 Prompt 适用于：历年高考真题、原创单项选择题、语法与词汇选择题、助词／连词填空题、提示词改写题、动词活用题、补助动词／授受动词／复合动词题、敬语与交际用语题、形容词／副词／量词／句型题、语言运用、单句完形、语篇完形、阅读理解、指代题、句子插入题、段落排序题、内容一致题等高考日语新题型。

## 一、任务参数

正式执行前，先读取并锁定以下参数；用户未填写时采用括号内默认值。

- 【项目名称】：__________
- 【任务模式】：从零制作／审核重做（默认：审核重做）
- 【题目状态】：全文冻结／仅历年真题冻结／原创题允许修订（默认：全文冻结）
- 【旧答案用途】：仅作待审核对象，不得作为正确性证明
- 【答案汇总】：需要／不需要（默认：需要）
- 【阅读翻译范围】：全文翻译／重点句翻译／仅答案解析（默认：全文翻译）
- 【交付模式】：纯净成品／红字批注标注版（默认：红字批注标注版）
- 【目录类型】：静态目录／不制作目录（默认：静态目录）
- 【输出格式】：Word／正文预览／Markdown（默认：Word）
- 【最终交付数量】：__________
- 【是否只提交最终成品】：是／否（默认：是）
- 【其他特殊要求】：__________

用户本轮明确提出的要求优先级最高。若用户明确说明“题目已经定档”“不考虑调整题目”，则题干、选项、提示词、下划线、题号、题序、年份、文章段落及空格位置全部冻结，不得以任何理由改写。

## 二、资料优先级与冲突处理

### 1. 资料优先级

按以下顺序判断依据：

1. 用户本轮明确要求；
2. 官方高考题目、官方答案及权威题源；
3. 用户指定的题目母本；
4. 用户提供的原始讲义、知识资料和术语母本；
5. 旧答案及解析文档。

原始讲义只作为知识、术语、接续、用法边界和易混辨析的依据，不得整段照搬；最终只提炼本题真正考查的一个核心用法。旧答案及旧解析只作为待审核对象、错误定位依据及修订差异依据，不能反过来证明答案正确。

### 2. 冲突阻断机制

当官方题源、题目母本、原始讲义和旧解析相互矛盾时，不得静默选择、不得自行拼凑、不得用常识替代原文。必须先确认：

- 当前项目的最高母本是哪一份；
- 题目是否冻结；
- 是否存在官方答案；
- 冲突是否由排版、文字提取、乱码或缺页造成；
- 能否通过句法、语义、语境和考点独立锁定唯一答案。

在问题尚未解决前，不得输出貌似确定的最终解析。资料不足时，应明确列出缺失项并请求补充；不得编造题干、选项、文章内容或所谓“标准答案”。

## 三、内容数据库重建

不要直接在旧 Word 中边读边修。先在内部为每一道题建立结构化记录，最终 Word 只能由该内容数据库生成。每题至少包含：

- 唯一复核编号；
- 章节、板块、题号、题型；
- 冻结题干；
- 选项或提示词；
- 横线前文本；
- 横线后已有文本；
- 独立求得的答案；
- 答案回填后的完整日文；
- 准确中文翻译；
- 唯一精确考点；
- 锁定答案的具体线索；
- 必要的变形过程；
- 错误选项或其他候选的排除理由；
- 阅读题原文定位证据；
- 对应原始资料依据；
- 修订类型：知识性／文字性／格式性；
- 最终复核状态；
- 版本专属的标音或无标音数据。

复核编号示例：CH03-POT-Q012。任何题号、答案、翻译、考点、解析、批注和答案汇总都必须通过该编号关联，防止串题、漏题和重复题。

## 四、总执行流程

逐题严格执行以下闭环：

题干 → 选项／提示词 → 独立作答 → 答案边界确认 → 完整日文 → 中文翻译 → 精确考点 → 解析 → 错误项排除 → 阅读证据或语篇依据 → 答案汇总同步 → 字体与版式处理 → 终审。

任何一项发生修改，其他各项必须重新核对。不得只改答案而不改翻译和解析，也不得只改解析而不核对答案汇总。

## 五、独立作答与反锚定

1. 正式判断答案时，先遮蔽旧答案和旧解析，从题干、选项、文章和权威资料出发独立作答。
2. 独立答案确定后，再与旧答案比较。
3. 旧答案即使“碰巧正确”，也必须重新确认完整日文、翻译、考点、判断理由和排除项。
4. 禁止“先看答案，再围绕答案反向编造题干含义或解析”。
5. 禁止把上一题、同板块其他题或原始讲义中的例句移植到本题解析中。

## 六、答案边界复核

答案不仅要正确，还必须只包含横线实际缺失的部分。

### 1. 无括号提示的填空题

横线原则上只能填写助词或明确的助词组合。不得把「な」「なさい」「て」「ろ」等孤立语法后缀或词尾碎片当作无提示填空答案。

### 2. 有括号提示的改写题

括号内应为完整名词、动词、一类形容词、二类形容词等。学生填写由提示词形成的完整、有效形式。

### 3. 横线后已有成分

必须保存并核对：

提示词原形 → 学生实际填写答案 → 题干后已有成分 → 回填后的完整形式。

- 题干后未给出「ます／ません／ました／ませんでした」等成分时，填写完整谓语；
- 题干已经给出「て／た／ば／そうだ／くれる／おく」等后续成分时，只填写横线所缺部分；
- 禁止答案与题干后缀重复；
- 禁止漏掉必要活用成分；
- 禁止为方便解析而改变横线边界。

最终必须满足：原题前段＋答案＋原题后段＝完整、自然且与母本一致的日文。

## 七、唯一答案四重验证

每一道题必须同时通过：

1. 【句法唯一】其他候选不能构成正确结构；
2. 【语义唯一】其他候选不能表达符合题意的意义；
3. 【语境唯一】题干或全文提供了决定时态、肯否、人物关系、指代或逻辑的有效线索；
4. 【考点唯一】横线或问题真正考查本板块知识。

不能只证明标准答案“可以”，还要证明高频候选“不可以”。

### 动词提示题

至少排查原形、ます形、ない形、た形、て形、可能形、被动形、使役形、使役被动形、意志形、命令形、ば形以及时态、肯否、简体、敬体的可能组合。

### 助词题

至少尝试「は、が、の、を、に、へ、で、と、から、まで、より、も、でも、しか」及相关复合助词，并说明高频干扰项为何不成立。

若题目允许修订且两个以上答案自然成立，必须补足语境或重写；若题目冻结，则必须依据官方答案、题源和上下文锁定解释口径，不得擅自改题。

## 八、完整日文与中文翻译

### 1. 完整日文

完整日文只能由冻结题干前段、最终答案和冻结题干后段自动合成。不得凭记忆重新复述题干，不得替换人物、时间、地点、数量、助词、语气或动作。

### 2. 中文翻译

根据审核后的完整日文重新翻译，逐项检查：

- 主语和人物关系；
- 动作对象；
- 时间、地点、数量；
- 时态、肯否；
- 被动、使役、授受；
- 敬语方向；
- 指代范围；
- 因果、转折、并列、递进、假定等逻辑。

不得加入原题中不存在的原因、目的、情绪、身份、事件结果或背景故事。翻译须自然，但不能牺牲原文信息准确性。

## 九、各题型解析规则

### 1. 语法／词汇／助词／副词／敬语单项选择题

统一格式：

答案字母 ☞ 准确中文翻译。
◆考点：考查「具体形式」在本题中的具体功能。
◆解析：先写正确项形成的完整结构及成立依据，再逐项说明其他选项不成立的具体原因。

错误项必须明确指出属于：接续错误、词义不符、时态不符、肯否不符、人称不符、敬语主体错误、语域不当、句末呼应不成立、逻辑关系错误或固定搭配不成立。禁止只写“其他选项不符合题意”。

### 2. 无提示助词／连词填空

重点检查固定搭配、自动词与他动词、移动经过与到达点、动作场所、被动施动者、使役对象、授受关系、引用内容、时间起止和复合助词。解析必须说明为什么该助词成立，以及至少两个高频候选为何不成立。

### 3. 提示词改写题

必须写清：

- 提示词词性；
- 动词类别或形容词类别；
- 目标活用；
- 具体变形步骤；
- 时态、肯否；
- 简体或礼貌体；
- 题干后已有成分；
- 学生实际填写范围；
- 题干中的锁定线索。

禁止只写“变为可能形”“变为过去式”等笼统说明。

### 4. 单句完形填空

同时判断横线前后接续、句内逻辑、时态、指代、固定搭配、语义自然度及本板块主考点。答案回填后必须朗读检查完整句。

### 5. 语篇完形／语言运用

不得把空格当作孤立单句。先梳理：

- 全文主题；
- 各段功能；
- 人物关系和说话人立场；
- 时间推进；
- 指示代词和省略成分；
- 连词承接；
- 新旧信息；
- 前后照应。

每个空的解析必须说明：本句在全文中的功能＋前文依据＋后文依据＋正确项的衔接作用＋错误项破坏了哪一种逻辑。

### 6. 阅读理解

先识别题型，再建立原文证据链。

- 细节理解：定位原句，说明题干／选项与原文的同义改写关系；
- 指代题：写清指示词、代词或省略成分的完整指代范围；
- 原因／目的题：区分直接原因、背景原因、结果和目的；
- 主旨题：概括各段功能和全文反复出现的语义核心；
- 标题题：标题必须覆盖全文，不能只对应局部例子；
- 推断题：区分原文事实与必要推论，禁止过度推断；
- 作者态度题：依据评价词、语气副词、转折和结论判断；
- 句子插入题：检查指代、连词、新旧信息、时间顺序和段落功能；
- 段落排序题：检查首次提出、指代回收、逻辑推进和时序；
- 内容一致题：逐项回原文核对，不能依靠常识猜测。

阅读题统一格式：

答案字母 ☞ 准确中文概括或翻译。
◆定位：第X段／原文关键词「……」。
◆考点：考查细节理解／指代判断／主旨概括／合理推断等具体能力。
◆解析：说明原文证据、正确项与原文的改写关系，并逐项排除干扰项。

常见阅读干扰项：主体偷换、对象偷换、时间范围变化、数量变化、原因结果倒置、条件结论倒置、局部扩大为全文、“可能”扩大为“一定”、原文未提及、作者观点与他人观点互换、把例子误当主旨。

### 7. 敬语与交际用语

解析必须明确：谁是动作主体、谁是尊敬对象、谁属于说话人一方、使用尊他／自谦／郑重中的哪一类、当前场景是否得体，以及错误选项是否存在对自己使用尊他语、对上级使用不当命令或敬语方向颠倒等问题。

### 8. 量词、数字、日期与读音题

说明计数对象、量词选择、数字与量词的组合、促音、浊音、半浊音、特殊日期／时间／人数读音及其他选项错误原因。不得以单字默认音读替代词语和语境判断。

### 9. 汉字假名互译题

确认词性、完整词形、送り仮名边界、时态和上下文语义。不能只凭单个汉字或假名机械替换。

### 10. 对话／听力文字稿题（如提供文字稿）

结合人物身份、说话目的、时间顺序、指代、省略和最终行动判断。若没有音频或完整文字稿，不得臆测语气和未出现的信息。

## 十、考点写法

考点统一使用：

考查「具体形式」＋本题中的具体功能。

合格示例：

- 考查「も」接数量词后表示数量强调的用法。
- 考查「まで」表示极端举例的用法。
- 考查补助动词「V-てある」表示人为动作结果状态的用法。
- 考查一类动词「読む」可能形过去否定简体的变形。
- 考查被动句中「に」提示动作实施者的用法。
- 考查他动词使役句中「に」提示使役对象的用法。
- 考查阅读中「その」指代前文已出现内容的用法。

禁止写：考查相关用法、考查动词变化、考查语法搭配、考查句子结构、考查前后衔接、考查助词与连接成分，或在一题中罗列该语法全部用法。

同一具体考点在全书中使用完全一致的术语；同一形式的不同功能不得混为一谈。

## 十一、解析写法

解析必须包含：

1. 填入后形成的完整结构；
2. 该结构在本句中的准确含义；
3. 决定答案的具体词语、句末形式或语篇线索；
4. 必要的变形过程；
5. 单选题错误选项的具体排除理由；
6. 阅读题的原文定位和改写关系；
7. 语篇题的全文逻辑依据。

禁止出现只有套话而无证据的表达，例如：

- “根据句意及后续结构……”；
- “与后续内容衔接……”；
- “该形式与主语、助词、时态及语气一致……”；
- “结合该表达的接续和语义，应使用……”；
- 只重复中文翻译，不解释判断依据；
- 罗列与本题无关的整章讲义知识。

## 十二、不合格旧解析的处理

采用“一票否决＋整题重做”。只要出现以下任意问题，原题的答案、翻译、考点和解析整块全部判定不合格：

- 答案错误；
- 答案虽正确但理由错误；
- 翻译与原题不对应；
- 解析引用另一道题；
- 擅自增加人物、时间、地点、动作或因果；
- 考点笼统或错判；
- 语篇题脱离全文；
- 阅读题没有原文证据；
- 错误项未逐项排除；
- 变形过程不完整；
- 答案边界多填或少填；
- 答案汇总与逐题答案不一致。

处理方式：删除原有“答案＋翻译＋考点＋解析”整块，依据冻结题目、文章和权威资料重新建立完整闭环。不得在幻觉内容上局部修补。

## 十三、反幻觉与可追溯机制

1. 解析中出现的人物、时间、地点、数量、动作、对象、肯否和因果，必须能在题目、文章或指定资料中找到依据。
2. 阅读题答案必须能够指向原文直接证据、前后逻辑或必要推论链。
3. 完整日文不得为了“更自然”而改写冻结题干。
4. 不得将常识、模型记忆或其他题目内容冒充本题依据。
5. 不确定时不得使用“可能是”“大概是”后直接交付成品；必须继续核查或阻断交付。
6. 每道题的答案、翻译、考点、解析、排除项、答案汇总和批注必须通过复核编号可追溯。

## 十四、答案汇总与静态目录

### 1. 答案汇总

答案汇总从最终逐题数据自动生成，不得手工维护第二套答案。必须保证：

答案汇总＝逐题答案＝完整日文中实际填入内容。

每10题或20题分组排版；组内题号连续；不得漏题、重题或错位。答案实际未改变时，不因解析重写而把答案汇总整段标红。

### 2. 静态目录

最终目录在全部内容和分页完成后生成并更新，随后转为稳定的静态可见目录。必须：

- 不显示「TOC \\o」等域代码；
- 目录页码与最终分页一致；
- 标题层级准确；
- 答案、考点和解析不得误设为标题；
- 可保留内部跳转，但可见文字不得包含字段指令。

## 十五、Word红字与批注

由【交付模式】决定是否保留修订标记。

### 1. 知识性修改

包括答案、翻译、考点、解析、错位、整题重构、多解或无解处理。

- 红字批注标注版：修改后的正确内容标红，并添加一条综合 Word 批注；
- 纯净成品：直接保留最终正确内容，不显示红字和批注。

批注模板：

【复核编号】CH03-POT-Q012
【问题】原解析与题干错位，且变形依据错误。
【处理】已依据定档题目重新核定答案并重写翻译、考点和解析。
【联动】答案汇总及相关引用已同步更新。

### 2. 非知识性修改

错别字、序号、标点、空格、语法符号统一、“考察”改为“考查”：标注版只标红实际改动文字，不加批注；纯净版直接修正。

### 3. 纯格式修改

字体、字号、语言属性、行距、段距、分页和标音结构：不标红、不加批注。

禁止使用会与标音结构交叉的跟踪修订节点。批注锚点不得插入日文单词、送り仮名或标音基底内部。

## 十六、字符清洗与异常空格

生成 Word 前后各执行一次字符审计，清除或阻断：

- Unicode 替代字符「�」；
- 缺字符方框和对象替代字符；
- 私用区及非法控制字符；
- 零宽空格、零宽不换行空格；
- 软连字符；
- 无语义的不间断空格；
- 单词内部异常制表符；
- 无意义全角空格；
- 连续普通空格；
- 多余手动换行；
- 裸露目录域代码；
- 隐藏删除文字和残留修订文本。

日文中的「ー」「っ／ッ」、拗音、浊音、半浊音、日文引号「」和中点「・」必须正确保留。正文字符缩放100%，字符间距0磅，位置为标准；不得使用分散对齐拉开日文汉字与假名。

## 十七、中日文字体双向审计

中文和日文不能只按“是否为汉字”判断，因为中文汉字和日文汉字共享字符编码。必须先按内容功能切分语言，再设置字体。

- 所有中文：宋体，11磅，语言属性 zh-CN；
- 所有日文，包括日汉字、平假名、片假名、日文标点及日文语境中的数字／字母：MS Mincho，11磅，语言属性 ja-JP；
- 中文题源信息使用宋体；
- 日文中的中文释义使用宋体；
- 中文解析中引用的日文词语、例句和语法形式使用MS Mincho；
- 中日文混排必须拆分为独立文本运行块；
- 四个字体槽 ascii、hAnsi、eastAsia、cs 均需明确设置，禁止依赖主题字体或自动回退。

必须双向检查：中文误用MS Mincho为0；平假名／片假名误用宋体为0；日汉字误用宋体为0；中文汉字误判为日文为0。

{VERSION_MODULE}

## 十九、版式与视觉检查

1. A4页面，页边距、页眉页脚、页码和标题样式统一；
2. 标题与下一段保持同页；“答案—考点—解析”尽量连续；
3. “◆考点”不得孤立在页尾；
4. 答案汇总的每组答案不得跨页拆散；
5. 表格行不跨页；长解析允许自然分页，不使用大量软换行硬撑版面；
6. 纯中文、答案汇总和标题不因标音需求全局放大行距；
7. 含日文标音的段落仅设置足够容纳上标的安全行距；
8. 不得出现乱码、方框、重叠、裁切、孤行、空白页、页码漂移或目录页码错误；
9. Word导航窗格不得出现括号读音、答案行或解析行；
10. 最终Word必须整本渲染为逐页图片，完成100%逐页视觉检查，不能只抽查首页和末页。

## 二十、全量内容终审

交付前逐题确认：

- 题目、选项、提示词、年份和题号符合【题目状态】；
- 题目没有遗漏、重复和错序；
- 答案独立成立且唯一；
- 答案边界准确；
- 回填完整日文无丢字、重复和倒序；
- 中文翻译与完整日文逐项对应；
- 考点具体、统一；
- 解析与本题完全对应；
- 单选题错误项逐项排除；
- 语篇题结合全文；
- 阅读题有清晰定位证据；
- 敬语主体和尊敬方向准确；
- 量词和读音依据语境确认；
- 答案汇总与逐题答案一致；
- Word批注、红字和复核编号对应；
- 字体、字号、语言属性、字符和分页合格。

## 二十一、硬性清零清单

最终交付前，下列项目必须全部为0：

- 冻结题目改动；
- 题目遗漏、重复、错位；
- 多解题、无解题；
- 答案错误或边界错误；
- 题目与解析错位；
- 翻译与完整日文不对应；
- 阅读答案无原文证据；
- 语篇题脱离全文；
- 考点笼统；
- 机械套话解析；
- 原题外虚构人物、时间、地点、动作和因果；
- 答案汇总不一致；
- 错误字符、乱码、方框；
- 异常空格、零宽字符和多余手动换行；
- 中文误用MS Mincho；
- 日文假名或日汉字误用宋体；
- 裸露TOC字段代码；
- 页面裁切、重叠、空白页和错页；
- 知识性修改漏红字或漏批注；
- {VERSION_ZERO_ITEMS}

## 二十二、最终输出要求

1. 严格按照用户要求的文件数量、文件名和格式交付；
2. 用户要求“只输出最终Word”时，不提交中间稿、台账、渲染图片或过程报告；
3. 用户要求压缩包时，压缩包内只放最终成品，不夹带临时文件；
4. 输出前检查文件能够正常打开、文件名准确、压缩包可解压、内部文件数量正确；
5. 只有全部终审项目通过后，才能声称“已完成并确认无误”；
6. 若终审未通过，不得用旧文件、临时文件或空壳文件冒充最终成品。

## 二十三、可直接套用的输出模板

### 1. 单项选择题

【题号】答案字母 ☞ 准确中文翻译。
◆考点：考查「具体形式」在本题中的具体功能。
◆解析：完整结构＋正确项依据＋题干线索＋A／B／C／D逐项排除。

### 2. 填空／提示词改写题

【题号】答案 ☞ 准确中文翻译。
◆考点：考查词性、具体活用、时态、肯否、语体及本题功能。
◆解析：提示词原形＋变形步骤＋题干已有后缀＋完整结构＋锁定线索。

### 3. 语篇完形题

【全文翻译】按【阅读翻译范围】执行。
【题号】答案 ☞ 本空所在句的准确翻译。
◆考点：考查具体词语／语法／连词在全文中的功能。
◆解析：本句功能＋前文依据＋后文依据＋指代／时态／人物立场＋错误项排除。

### 4. 阅读理解题

【题号】答案字母 ☞ 准确中文概括。
◆定位：第X段「原文关键词或关键句」。
◆考点：考查细节理解／指代判断／主旨概括／合理推断／作者态度等。
◆解析：原文证据＋正确项改写关系＋错误项类型及排除理由。

## 二十四、开始执行

读取全部文件后，先锁定任务参数和题目状态，再按照“内容数据库重建＋字符清洗＋答案边界复核＋双向字体审计＋静态目录＋全量内容终审”的流程连续执行。除非资料缺失导致无法唯一判断，否则不要中途提交半成品。正式完成并通过全部终审后，只按用户要求交付最终成品。
'''

RUBY_MODULE = r'''
## 十八、含标音版专项规则

本版本必须使用Word原生Ruby结构对指定日文内容进行单字逐字标音。

### 1. 字体与字号

- 所有日文正文，包括日汉字、平假名、片假名和日文标点：MS Mincho，11磅；
- 所有中文：宋体，11磅；
- 所有日文上标假名：MS Mincho，5.5磅；
- 纯字体、字号、语言属性和Ruby结构调整不标红、不加批注。

### 2. 允许标音的区域

- 逐题详细答案中的日文答案词语；
- 答案回填后的完整日文；
- 解析中明确引用的日文词语、结构和例句；
- 日文选项；
- 日文语篇和阅读原文；
- 日文人名、地名、专有名词。

### 3. 永久禁止标音的区域

- 封面、静态目录和所有标题；
- 页眉、页脚和页码；
- Word批注；
- 中文翻译、中文考点和中文解析；
- 单独提取的答案速查汇总；
- 题号、答案字母；
- 日文中的中文释义。

特别强调：答案汇总专供学生快速核对，即使包含日文汉字，也不得追加Ruby。

### 4. 单字Ruby

一个日汉字对应一个Ruby基底，不得把多个汉字合并成整词标音。送り仮名保留在正文基线上，不进入Ruby基底。

- 「日本語」：日〔に〕＋本〔ほん〕＋語〔ご〕；
- 「お金」：お＋金〔かね〕，不能把「金」标成「きん」；
- 「諦めない」：諦〔あきら〕＋めない；
- 「心」：心〔こころ〕；
- 「行きます」：行〔い〕＋きます；
- 「1日」：根据当前语境确认；当题目指定为「いちにち」时，「日」标「にち」，不得全局机械替换。

### 5. 读音确认

先识别完整词语、词性、句义和送り仮名边界，再判断音读／训读并拆成单字Ruby。同一汉字不得执行全局读音替换，例如「お金」的「金」读「かね」，「金額」的「金」读「きん」。日期、量词、人名、地名、熟字训、当て字、连浊、促音、浊音、半浊音、敬语特殊词和同形异读词必须逐处人工确认。

不能百分之百确认的读音进入异常清单，解决前禁止交付；不得使用“最常见读音”作为自动回退。

### 6. Word原生结构

- 使用Word原生<w:ruby>结构；
- <w:rubyBase>中只允许一个日汉字；
- <w:rt>中的假名不能为空；
- Ruby基底为MS Mincho 11磅；
- Ruby上标为MS Mincho 5.5磅；
- 中文汉字不得进入Ruby节点；
- 不得用普通上标、括号、文本框或域代码模拟Ruby。

### 7. Ruby前后正文一致性

生成Ruby前保存无标音日文，生成后提取普通正文和Ruby基底重新组合。必须满足：生成前日文＝生成后正文基底。丢字、增字、倒序、重复、异常空格均为0。

### 8. 含标音版专项审计

- 中文区域Ruby：0；
- 答案汇总Ruby：0；
- 标题、目录、页眉页脚Ruby：0；
- 日文汉字漏标：0；
- 错误读音：0；
- 整词合并Ruby：0；
- 送り仮名误入Ruby基底：0；
- 空Ruby：0；
- Ruby字号或字体错误：0；
- Ruby造成正文变化：0；
- Ruby重叠、裁切和行距异常：0。
'''

PLAIN_MODULE = r'''
## 十八、不含标音版专项规则

本版本全文禁止系统新增任何日文读音标注。

### 1. 禁止形式

- 不生成Word原生Ruby节点；
- 不使用普通上标模拟读音；
- 不在日汉字后自动添加括号读音；
- 不使用文本框、域代码或隐藏文字模拟标音；
- 不因删除标音而丢失、重复或改变正文。

### 2. 原题自带标音

设置参数【保留原题自带标音】是／否。默认规则：原题本身已有且属于冻结题目内容的标音原样保留；系统不得新增任何读音；解析中新引用的日文词语不额外标音。若用户要求“全文无标音”，则连原题已有的辅助标音也按用户要求处理，但不得修改题目正文。

### 3. 字体规则不变

不含标音不等于不区分语言：

- 中文：宋体11磅，zh-CN；
- 日文汉字、平假名、片假名、日文标点：MS Mincho 11磅，ja-JP；
- 中文解析中引用的日文语法形式：MS Mincho；
- 日文中的中文释义：宋体；
- 中日文混排继续拆分为独立运行块。

### 4. 与含标音版的内容一致性

若同一项目同时制作含标音版和不含标音版，两版必须来自同一内容数据库。题目数量、答案、翻译、考点、解析、错误项排除、阅读定位、章节、题号、红字和批注逻辑完全一致；唯一差异只能是Ruby及相应的安全行距。

### 5. 不含标音版专项审计

- Word原生Ruby节点：0；
- 普通上标模拟读音：0；
- 系统新增括号读音：0；
- 文本框或域代码模拟标音：0；
- 标音残留造成的异常行距：0；
- 删除标音导致的丢字、增字或异常空格：0。
'''


def build_text(ruby: bool) -> str:
    module = RUBY_MODULE if ruby else PLAIN_MODULE
    zero_items = (
        "中文误加Ruby；答案汇总误加Ruby；标题或目录误加Ruby；日文汉字漏标或错标；整词Ruby；Ruby造成正文变化"
        if ruby
        else "Ruby节点、模拟上标读音、系统新增括号读音及标音残留"
    )
    return COMMON_PROMPT.replace("{VERSION_MODULE}", module).replace("{VERSION_ZERO_ITEMS}", zero_items)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, font_name: str, size: float = 11.0, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)


def add_lang(run, lang: str) -> None:
    rpr = run._element.get_or_add_rPr()
    lang_el = rpr.find(qn("w:lang"))
    if lang_el is None:
        lang_el = OxmlElement("w:lang")
        rpr.append(lang_el)
    lang_el.set(qn("w:val"), lang)
    lang_el.set(qn("w:eastAsia"), lang)


def contains_kana(text: str) -> bool:
    return bool(re.search(r"[\u3040-\u30ff]", text))


def add_mixed_text(paragraph, text: str, *, bold: bool = False, size: float = 11.0, color: RGBColor | None = None) -> None:
    # Corner-quoted spans are Japanese terms/examples. Outside them, kana runs are Japanese;
    # Chinese explanatory text remains SimSun. Full Japanese example lines are handled as one Japanese run.
    if contains_kana(text) and (text.startswith("JP|") or text.startswith("【日文】")):
        clean = text.split("|", 1)[1] if text.startswith("JP|") else text.replace("【日文】", "", 1)
        run = paragraph.add_run(clean)
        set_run_font(run, "MS Mincho", size, bold, color)
        add_lang(run, "ja-JP")
        return

    parts = re.split(r"(「[^」]*」|『[^』]*』)", text)
    for part in parts:
        if not part:
            continue
        if (part.startswith("「") and part.endswith("」")) or (part.startswith("『") and part.endswith("』")):
            run = paragraph.add_run(part)
            set_run_font(run, "MS Mincho", size, bold, color)
            add_lang(run, "ja-JP")
            continue
        # Split kana/formula sequences from Chinese prose.
        subparts = re.split(r"([A-Za-z0-9＋＋\-・／～〜\u3040-\u30ff]+)", part)
        for sub in subparts:
            if not sub:
                continue
            if contains_kana(sub):
                run = paragraph.add_run(sub)
                set_run_font(run, "MS Mincho", size, bold, color)
                add_lang(run, "ja-JP")
            else:
                run = paragraph.add_run(sub)
                set_run_font(run, "宋体", size, bold, color)
                add_lang(run, "zh-CN")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(4)

    for name, size in (("Title", 18), ("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11.5)):
        style = doc.styles[name]
        style.font.name = "宋体"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)

    # Footer page number.
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run("第 ")
    set_run_font(r1, "宋体", 9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)
    r2 = p.add_run(" 页")
    set_run_font(r2, "宋体", 9)


def add_cover(doc: Document, title: str, ruby: bool) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, "宋体", 20, True)
    p.space_after = Pt(12)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("专门适配高考日语新题型｜单选・阅读・语言运用・语篇完形・提示词改写等")
    set_run_font(r2, "宋体", 12, False, RGBColor(89, 89, 89))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    variant = "Word原生单字Ruby输出模式" if ruby else "纯日文无新增标音输出模式"
    r3 = p3.add_run(variant)
    set_run_font(r3, "MS Mincho" if ruby else "宋体", 11, False, RGBColor(89, 89, 89))

    doc.add_paragraph()
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    box.columns[0].width = Cm(15.5)
    cell = box.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "F2F2F2")
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note = "本文件为可直接复制使用的最终Prompt。内容以同一出版级解析母体为基础，仅输出渲染规则不同。"
    add_mixed_text(cp, note, size=10.5)

    doc.add_page_break()


def add_usage(doc: Document, ruby: bool) -> None:
    h = doc.add_paragraph(style="Heading 1")
    add_mixed_text(h, "使用说明", bold=True, size=15)
    items = [
        "将本Prompt与题目文档、旧答案解析、官方题源及原始讲义一并提交。",
        "先填写“任务参数”，尤其确认题目是否冻结、交付模式和阅读翻译范围。",
        "审核重做模式下，旧答案仅作待审核对象；不合格解析按整题闭环重做。",
        "最终交付前必须完成内容审校、字符清洗、答案边界复核、字体双向审计、静态目录和整本视觉检查。",
        ("本版本启用Word原生单字Ruby；答案汇总、标题、目录和中文区域永久禁用Ruby。" if ruby else "本版本禁止系统新增Ruby、普通上标或括号读音；中日文字体仍须严格拆分。"),
    ]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_mixed_text(p, item)
    doc.add_paragraph()
    h2 = doc.add_paragraph(style="Heading 1")
    add_mixed_text(h2, "可直接复制使用的最终 Prompt", bold=True, size=15)


def add_prompt_body(doc: Document, text: str) -> None:
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_mixed_text(p, line[4:], bold=True, size=11.5)
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_mixed_text(p, line[3:], bold=True, size=13)
        elif line.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            add_mixed_text(p, line[2:], bold=True, size=15)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_mixed_text(p, line[2:])
        else:
            p = doc.add_paragraph()
            if re.match(r"^\d+\. ", line):
                p.paragraph_format.left_indent = Cm(0.2)
            add_mixed_text(p, line)


def build_doc(path: Path, title: str, ruby: bool) -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc, title, ruby)
    add_usage(doc, ruby)
    add_prompt_body(doc, build_text(ruby))
    # Core properties.
    doc.core_properties.title = title
    doc.core_properties.subject = "高考日语新题型配套答案及解析出版级制作与审校Prompt"
    doc.core_properties.author = "阿狸老师日语课堂"
    doc.core_properties.keywords = "高考日语, 新题型, 答案解析, 阅读理解, 完形填空, Ruby"
    doc.save(path)


def validate_docx(path: Path, ruby: bool) -> list[str]:
    problems: list[str] = []
    if not path.exists() or path.stat().st_size < 20000:
        problems.append(f"文件不存在或过小: {path}")
        return problems
    if not zipfile.is_zipfile(path):
        problems.append(f"不是有效DOCX压缩结构: {path}")
        return problems
    with zipfile.ZipFile(path, "r") as zf:
        bad = zf.testzip()
        if bad:
            problems.append(f"DOCX内部损坏: {bad}")
        names = set(zf.namelist())
        required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
        if not required.issubset(names):
            problems.append(f"DOCX缺少核心部件: {required - names}")
        xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
        for must in ("高考日语新题型", "内容数据库重建", "答案边界复核", "阅读理解", "双向字体审计", "静态目录", "全量内容终审"):
            if must not in xml:
                problems.append(f"缺少核心内容: {must}")
        if "TOC \\o" in xml:
            problems.append("存在裸露TOC域代码")
        if "\ufffd" in xml:
            problems.append("存在Unicode替代字符")
        if ruby:
            for must in ("Word原生Ruby", "答案汇总Ruby", "MS Mincho", "5.5磅", "お金"):
                if must not in xml:
                    problems.append(f"含标音版缺少专项内容: {must}")
        else:
            for must in ("全文禁止系统新增", "Word原生Ruby节点：0", "普通上标模拟读音：0"):
                if must not in xml:
                    problems.append(f"不含标音版缺少专项内容: {must}")
    return problems


def main() -> None:
    ruby_path = OUT_DIR / f"{TITLE_RUBY}.docx"
    plain_path = OUT_DIR / f"{TITLE_PLAIN}.docx"
    build_doc(ruby_path, TITLE_RUBY, True)
    build_doc(plain_path, TITLE_PLAIN, False)

    all_problems = validate_docx(ruby_path, True) + validate_docx(plain_path, False)
    if all_problems:
        raise SystemExit("\n".join(all_problems))

    # A manifest is printed to logs only; it is not included in the delivered artifact.
    print("Generated and validated:")
    for p in (ruby_path, plain_path):
        print(f"- {p.name}: {p.stat().st_size} bytes")


if __name__ == "__main__":
    main()
