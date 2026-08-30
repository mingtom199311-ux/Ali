#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
REPO_ROOT = ROOT.parent
OUTPUT_DIR = REPO_ROOT / "generated_prompt_docs_v2"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

WITH_RUBY_NAME = "Gaokao_Japanese_New_Format_Analysis_Prompt_With_Ruby.docx"
NO_RUBY_NAME = "Gaokao_Japanese_New_Format_Analysis_Prompt_No_Ruby.docx"
REPORT_NAME = "validation_report.json"

COMMON = (ROOT / "common_template.md").read_text(encoding="utf-8")
RUBY_RULES = (ROOT / "ruby_rules.md").read_text(encoding="utf-8")
NO_RUBY_RULES = (ROOT / "no_ruby_rules.md").read_text(encoding="utf-8")

JP_FONT = "MS Mincho"
CN_FONT = "宋体"

JP_TOKEN_RE = re.compile(
    r"(「[^」]*」|『[^』]*』|"
    r"(?:[A-Za-zＡ-Ｚａ-ｚ]+[-＋]?[ぁ-んァ-ヶー]+(?:[-＋／・]?[ぁ-んァ-ヶー]+)*)|"
    r"(?:[ぁ-んァ-ヶー]{2,}))"
)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, font_name: str, size: float = 11, bold: bool | None = None,
                 color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font_name)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    locale = "ja-JP" if font_name == JP_FONT else "zh-CN"
    lang.set(qn("w:val"), locale)
    lang.set(qn("w:eastAsia"), locale)


def add_mixed_text(paragraph, text: str, size: float = 11, bold: bool = False,
                   color: RGBColor | None = None) -> None:
    pos = 0
    for match in JP_TOKEN_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, CN_FONT, size, bold, color)
        run = paragraph.add_run(match.group(0))
        set_run_font(run, JP_FONT, size, bold, color)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, CN_FONT, size, bold, color)


def set_paragraph_spacing(paragraph, before=0, after=4, line=1.2) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = line
    fmt.widow_control = True


def set_keep_with_next(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if value and node is None:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    elif not value and node is not None:
        p_pr.remove(node)


def add_horizontal_rule(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B7C9E2")
    p_bdr.append(bottom)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, CN_FONT, 9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()._r
    r.append(fld_char1)
    r.append(instr_text)
    r.append(fld_char2)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, CN_FONT, 9)


def configure_document(doc: Document, title: str) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = CN_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    heading_specs = {
        "Title": (18, RGBColor(31, 78, 121)),
        "Subtitle": (11, RGBColor(89, 89, 89)),
        "Heading 1": (14, RGBColor(31, 78, 121)),
        "Heading 2": (12, RGBColor(47, 84, 150)),
        "Heading 3": (11, RGBColor(68, 68, 68)),
    }
    for style_name, (size, color) in heading_specs.items():
        style = styles[style_name]
        style.font.name = CN_FONT
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    for section in doc.sections:
        header_p = section.header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_mixed_text(header_p, "阿狸老师｜高考日语新题型配套解析最终Prompt", 8.5, False,
                       RGBColor(128, 128, 128))
        add_page_number(section.footer.paragraphs[0])

    doc.core_properties.title = title
    doc.core_properties.subject = "高考日语新题型配套解析出版级可复用Prompt"
    doc.core_properties.author = "阿狸老师日语课堂"
    doc.core_properties.keywords = "高考日语, 解析Prompt, 阅读理解, 完形填空, Ruby, Word审校"


def add_cover(doc: Document, title: str, version_name: str) -> None:
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_mixed_text(p, title, 18, True, RGBColor(31, 78, 121))
    set_paragraph_spacing(p, after=10, line=1.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_mixed_text(p, "出版级｜单选·填空·改写·完形·语篇·阅读·敬语·量词·句型", 11, False,
                   RGBColor(89, 89, 89))
    set_paragraph_spacing(p, after=16, line=1.0)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("版本", version_name),
        ("适用模式", "从零制作／审核重做；题目冻结／允许修订；纯净版／红字批注版"),
        ("核心闭环", "题干→独立答案→完整日文→中文翻译→精确考点→解析→排除项→答案汇总"),
        ("复制范围", "从“【最终可复制Prompt开始】”复制至“【最终可复制Prompt结束】”"),
    ]
    for r, (left, right) in enumerate(rows):
        for c in range(2):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if c == 0:
                set_cell_shading(cell, "D9EAF7")
        p1 = table.cell(r, 0).paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_mixed_text(p1, left, 10.5, True)
        p2 = table.cell(r, 1).paragraphs[0]
        add_mixed_text(p2, right, 10.5, False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=14, after=0)
    add_mixed_text(p, "编制依据：已验收通过的解析制作规则与字体／标音规则", 9.5, False,
                   RGBColor(100, 100, 100))

    doc.add_page_break()


def add_prompt_content(doc: Document, markdown_text: str) -> None:
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue
        if line.strip() == "---":
            p = doc.add_paragraph()
            add_horizontal_rule(p)
            set_paragraph_spacing(p, before=2, after=5)
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            add_mixed_text(p, line[2:].strip(), 14, True, RGBColor(31, 78, 121))
            set_paragraph_spacing(p, before=8, after=6, line=1.1)
            set_keep_with_next(p)
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_mixed_text(p, line[3:].strip(), 12, True, RGBColor(47, 84, 150))
            set_paragraph_spacing(p, before=8, after=4, line=1.1)
            set_keep_with_next(p)
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_mixed_text(p, line[4:].strip(), 11, True, RGBColor(68, 68, 68))
            set_paragraph_spacing(p, before=6, after=3, line=1.1)
            set_keep_with_next(p)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_mixed_text(p, line[2:].strip(), 11)
            p.paragraph_format.left_indent = Cm(0.65)
            p.paragraph_format.first_line_indent = Cm(-0.25)
            set_paragraph_spacing(p, after=2, line=1.18)
            continue
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.right_indent = Cm(0.5)
            set_cell = False
            add_mixed_text(p, line[2:].strip(), 10.5, False, RGBColor(64, 64, 64))
            set_paragraph_spacing(p, after=3, line=1.15)
            continue
        p = doc.add_paragraph()
        if line.startswith("【最终可复制Prompt"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_mixed_text(p, line, 11, True, RGBColor(192, 0, 0))
            set_paragraph_spacing(p, before=4, after=8, line=1.0)
            continue
        add_mixed_text(p, line, 11)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(p, after=3, line=1.22)


def build_prompt(version_name: str, section_title: str, version_rules: str,
                 filename: str) -> Path:
    prompt = COMMON.replace("{{VERSION_NAME}}", version_name)
    prompt = prompt.replace("{{VERSION_SECTION_TITLE}}", section_title)
    prompt = prompt.replace("{{VERSION_RULES}}", version_rules.strip())

    title = f"高考日语新题型配套解析最终Prompt（{version_name}）"
    doc = Document()
    configure_document(doc, title)
    add_cover(doc, title, version_name)
    add_prompt_content(doc, prompt)

    path = OUTPUT_DIR / filename
    doc.save(path)
    return path


def extract_doc_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def structural_validate(path: Path, version_name: str) -> dict:
    required = [
        "【最终可复制Prompt开始】",
        "【最终可复制Prompt结束】",
        "题干 →",
        "答案边界专项审核",
        "唯一答案四重验证",
        "阅读理解专项规则",
        "反幻觉与实体一致性",
        "字符清洗、中日文字体、静态目录与版式",
        "共同终审清零清单",
        "正式任务输入区",
        version_name,
    ]
    forbidden = ["{{VERSION", "�", "\u200b", "TOC \\o"]

    with zipfile.ZipFile(path, "r") as zf:
        bad = zf.testzip()
        if bad:
            raise AssertionError(f"Corrupt member {bad} in {path.name}")
        document_xml = zf.read("word/document.xml").decode("utf-8")
        if "w:ruby" in document_xml:
            raise AssertionError("Prompt instruction document itself must not contain active Ruby nodes")
        if CN_FONT not in document_xml or JP_FONT not in document_xml:
            raise AssertionError("Expected Chinese/Japanese font declarations missing")

    text = extract_doc_text(path)
    for item in required:
        if item not in text:
            raise AssertionError(f"Missing required marker: {item}")
    for item in forbidden:
        if item in text:
            raise AssertionError(f"Forbidden text found: {item}")
    if len(text) < 20000:
        raise AssertionError(f"Prompt content unexpectedly short: {len(text)}")
    if path.stat().st_size < 30000:
        raise AssertionError(f"DOCX unexpectedly small: {path.stat().st_size}")

    return {
        "filename": path.name,
        "bytes": path.stat().st_size,
        "characters": len(text),
        "required_markers": "pass",
        "zip_integrity": "pass",
        "font_declarations": "pass",
        "active_ruby_nodes_in_prompt_document": 0,
    }


def render_validate(paths: Iterable[Path]) -> dict:
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        return {"status": "skipped", "reason": "LibreOffice not installed"}

    import fitz  # PyMuPDF, imported only when rendering is requested

    render_report = {"status": "pass", "documents": []}
    with tempfile.TemporaryDirectory(prefix="prompt_render_") as tmp:
        tmpdir = Path(tmp)
        for path in paths:
            cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(tmpdir), str(path)]
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
            if proc.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed for {path.name}: {proc.stderr}")
            pdf_path = tmpdir / f"{path.stem}.pdf"
            if not pdf_path.exists():
                raise AssertionError(f"Rendered PDF missing for {path.name}")
            pdf = fitz.open(pdf_path)
            if pdf.page_count < 8:
                raise AssertionError(f"Unexpectedly low page count for {path.name}: {pdf.page_count}")
            page_stats = []
            for idx, page in enumerate(pdf):
                text = page.get_text("text").strip()
                if len(text) < 5:
                    raise AssertionError(f"Blank or nearly blank rendered page {idx + 1} in {path.name}")
                pix = page.get_pixmap(matrix=fitz.Matrix(0.7, 0.7), alpha=False, colorspace=fitz.csGRAY)
                samples = pix.samples
                dark = sum(1 for b in samples if b < 245)
                ratio = dark / max(1, len(samples))
                if ratio < 0.0003:
                    raise AssertionError(f"Visually near-empty page {idx + 1} in {path.name}")
                page_stats.append({"page": idx + 1, "text_chars": len(text), "ink_ratio": round(ratio, 6)})
            render_report["documents"].append({
                "filename": path.name,
                "pages": pdf.page_count,
                "all_pages_nonblank": True,
                "min_text_chars": min(x["text_chars"] for x in page_stats),
                "min_ink_ratio": min(x["ink_ratio"] for x in page_stats),
            })
            pdf.close()
    return render_report


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--render", action="store_true", help="also render all pages through LibreOffice")
    args = parser.parse_args()

    with_ruby = build_prompt(
        version_name="含标音版",
        section_title="含标音版｜单字Ruby专项规则",
        version_rules=RUBY_RULES,
        filename=WITH_RUBY_NAME,
    )
    no_ruby = build_prompt(
        version_name="不含标音版",
        section_title="不含标音版｜全文禁新增标音专项规则",
        version_rules=NO_RUBY_RULES,
        filename=NO_RUBY_NAME,
    )

    report = {
        "status": "pass",
        "documents": [
            structural_validate(with_ruby, "含标音版"),
            structural_validate(no_ruby, "不含标音版"),
        ],
    }
    if args.render:
        report["render_validation"] = render_validate([with_ruby, no_ruby])
        if report["render_validation"].get("status") != "pass":
            raise AssertionError("Rendering was requested but did not complete")

    report_path = OUTPUT_DIR / REPORT_NAME
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
