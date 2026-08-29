from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path
from typing import Iterable, List, Sequence, Tuple

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "output"
OUTPUT.mkdir(parents=True, exist_ok=True)

COMMON_FILES = [
    ROOT / "01_front_and_principles.md",
    ROOT / "02_workflow_and_formats.md",
    ROOT / "03_question_modules.md",
    ROOT / "04_word_and_audit.md",
]
FINAL_FILE = ROOT / "05_final_checklist.md"

VERSIONS = [
    {
        "key": "ruby",
        "name": "含单字ルビ标音版",
        "title": "高考日语新题型配套解析最终Prompt",
        "subtitle": "含单字ルビ标音版",
        "module": ROOT / "module_ruby.md",
        "filename": "高考日语新题型配套解析最终Prompt（含单字ルビ标音版）.docx",
        "positioning": "适用于需要Word原生单字ルビ的配套答案与解析制作。答案速查汇总、标题、目录和中文内容禁止标音。",
    },
    {
        "key": "no_ruby",
        "name": "不含标音版",
        "title": "高考日语新题型配套解析最终Prompt",
        "subtitle": "不含标音版",
        "module": ROOT / "module_no_ruby.md",
        "filename": "高考日语新题型配套解析最终Prompt（不含标音版）.docx",
        "positioning": "适用于不新增任何Ruby、上标假名或括号读音的配套答案与解析制作。内容审核标准与含标音版完全一致。",
    },
]

CN_FONT = "宋体"
JP_FONT = "MS Mincho"
BODY_SIZE = Pt(11)
RUBY_SIZE = Pt(5.5)
ACCENT = RGBColor(31, 102, 111)
DARK = RGBColor(38, 47, 52)
MUTED = RGBColor(92, 104, 110)
LIGHT_FILL = "EAF2F3"
PALE_FILL = "F6F8F8"
BORDER = "AAB7BA"

HIRAGANA_KATAKANA_RE = re.compile(r"[\u3040-\u30ff\uff66-\uff9f]")
JP_QUOTE_RE = re.compile(r"(「[^」]*」|『[^』]*』)")
INLINE_TOKEN_RE = re.compile(r"(\*\*[^*]+\*\*|「[^」]*」|『[^』]*』|`[^`]+`)")
TABLE_SEPARATOR_RE = re.compile(r"^\|(?:\s*:?-+:?\s*\|)+$")
NUMBERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ("sz", "val", "color", "space"):
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_run_font(run, font: str, size: Pt = BODY_SIZE, bold: bool | None = None,
                 italic: bool | None = None, color: RGBColor | None = None,
                 lang: str | None = None) -> None:
    run.font.name = font
    run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font)
    if lang:
        lang_el = rpr.find(qn("w:lang"))
        if lang_el is None:
            lang_el = OxmlElement("w:lang")
            rpr.append(lang_el)
        lang_el.set(qn("w:val"), lang)
        lang_el.set(qn("w:eastAsia"), lang)


def looks_japanese(text: str) -> bool:
    return bool(HIRAGANA_KATAKANA_RE.search(text))


def add_text_runs(paragraph, text: str, *, default_bold: bool = False,
                  default_italic: bool = False, default_color: RGBColor | None = None,
                  force_japanese: bool = False) -> None:
    """Add text with Japanese quotes/kana in MS Mincho and Chinese in SimSun.

    Markdown bold and inline code markers are stripped while retaining visual emphasis.
    """
    if not text:
        return
    parts = INLINE_TOKEN_RE.split(text)
    for part in parts:
        if not part:
            continue
        bold = default_bold
        italic = default_italic
        color = default_color
        content = part
        if part.startswith("**") and part.endswith("**"):
            bold = True
            content = part[2:-2]
        elif part.startswith("`") and part.endswith("`"):
            content = part[1:-1]
            color = ACCENT
        is_jp = force_japanese or bool(JP_QUOTE_RE.fullmatch(content)) or looks_japanese(content)
        font = JP_FONT if is_jp else CN_FONT
        lang = "ja-JP" if is_jp else "zh-CN"
        run = paragraph.add_run(content)
        set_run_font(run, font, BODY_SIZE, bold=bold, italic=italic, color=color, lang=lang)


def configure_document(doc: Document, meta: dict) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.15)
    sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = CN_FONT
    normal.font.size = BODY_SIZE
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.widow_control = True

    for name, size, color in [
        ("Title", 22, ACCENT),
        ("Subtitle", 12.5, MUTED),
        ("Heading 1", 15.5, ACCENT),
        ("Heading 2", 13.5, ACCENT),
        ("Heading 3", 11.5, DARK),
        ("Heading 4", 11, DARK),
    ]:
        style = styles[name]
        style.font.name = CN_FONT
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)

    styles["List Bullet"].font.name = CN_FONT
    styles["List Bullet"].font.size = BODY_SIZE
    styles["List Bullet"]._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    styles["List Number"].font.name = CN_FONT
    styles["List Number"].font.size = BODY_SIZE
    styles["List Number"]._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    props = doc.core_properties
    props.title = f"{meta['title']}（{meta['subtitle']}）"
    props.subject = "高考日语新题型配套答案与解析出版级可复用Prompt"
    props.author = "阿狸老师日语课堂"
    props.keywords = "高考日语,答案解析,新题型,单选,阅读,语篇,完形,Word,Ruby"
    props.comments = "依据已验收规则整理的最终可复用Prompt。"


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, CN_FONT, Pt(9), color=MUTED, lang="zh-CN")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), CN_FONT)
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rpr.append(sz)
    r.append(rpr)
    r.append(fld_char1)
    r.append(instr_text)
    r.append(fld_char2)
    paragraph._p.append(r)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, CN_FONT, Pt(9), color=MUTED, lang="zh-CN")


def add_header_footer(doc: Document, meta: dict) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"阿狸老师日语课堂｜{meta['subtitle']}")
        set_run_font(run, CN_FONT, Pt(8.5), color=MUTED, lang="zh-CN")
        footer = section.footer
        add_page_field(footer.paragraphs[0])


def add_cover(doc: Document, meta: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta["title"])
    set_run_font(r, CN_FONT, Pt(22), bold=True, color=ACCENT, lang="zh-CN")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(meta["subtitle"])
    set_run_font(r2, CN_FONT, Pt(15), bold=True, color=DARK, lang="zh-CN")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(15)
    add_text_runs(p3, "单选｜填空｜提示词改写｜语篇完形｜阅读理解｜敬语交际｜量词句型", default_color=MUTED)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(15.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=180, bottom=180, start=220, end=220)
    set_cell_border(cell, top={"val": "single", "sz": "8", "color": "8FB2B6"},
                    bottom={"val": "single", "sz": "8", "color": "8FB2B6"},
                    start={"val": "single", "sz": "8", "color": "8FB2B6"},
                    end={"val": "single", "sz": "8", "color": "8FB2B6"})
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_runs(cp, meta["positioning"], default_color=DARK)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(42)
    add_text_runs(p4, "最终可复制版｜内容数据库重建＋答案边界复核＋反幻觉审校＋双向字体审计", default_color=MUTED)
    doc.add_page_break()


def extract_headings(markdown: str) -> List[Tuple[int, str]]:
    out: List[Tuple[int, str]] = []
    for line in markdown.splitlines():
        m = re.match(r"^(#{2,3})\s+(.+)$", line.strip())
        if m:
            out.append((len(m.group(1)), m.group(2).strip()))
    return out


def add_usage_and_navigation(doc: Document, meta: dict, markdown: str) -> None:
    h = doc.add_heading("使用说明", level=1)
    add_text_runs(doc.add_paragraph(), "本文件中的完整Prompt可直接复制使用。执行前只需填写参数区，并同时上传题目母本、原始讲义、旧答案解析及权威题源。")

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    items = [
        ("第一步", "填写任务模式、题目冻结状态、交付模式和最终文件名。"),
        ("第二步", "上传题目母本、NO.1—NO.8等知识资料及旧答案解析。"),
        ("第三步", "从“可直接复制使用的完整Prompt”开始复制至“最终执行命令”结束。"),
        ("双版本原则", "含标音版与不含标音版共用同一内容数据库，只允许在标音渲染模块上不同。"),
    ]
    for row, (a, b) in zip(table.rows, items):
        row.cells[0].width = Cm(3.0)
        row.cells[1].width = Cm(13.5)
        set_cell_shading(row.cells[0], LIGHT_FILL)
        for c in row.cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(c)
        p0 = row.cells[0].paragraphs[0]
        add_text_runs(p0, a, default_bold=True, default_color=ACCENT)
        p1 = row.cells[1].paragraphs[0]
        add_text_runs(p1, b)

    doc.add_heading("内容导航", level=1)
    for level, title in extract_headings(markdown):
        if level == 2:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.25)
            p.paragraph_format.space_after = Pt(1)
            add_text_runs(p, title, default_bold=True, default_color=DARK)
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("可直接复制使用的完整 Prompt")
    set_run_font(r, CN_FONT, Pt(17), bold=True, color=ACCENT, lang="zh-CN")

    info = doc.add_table(rows=1, cols=1)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = info.cell(0, 0)
    set_cell_shading(cell, PALE_FILL)
    set_cell_margins(cell, top=130, bottom=130, start=180, end=180)
    ip = cell.paragraphs[0]
    add_text_runs(ip, "复制范围：从下一段“你现在是一名……”开始，直至文末“最终执行命令”结束。")


def add_table_from_lines(doc: Document, lines: Sequence[str]) -> None:
    parsed = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        parsed.append(cells)
    if len(parsed) >= 2 and all(re.fullmatch(r":?-+:?", x.replace(" ", "")) for x in parsed[1]):
        parsed.pop(1)
    if not parsed:
        return
    cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row_data in enumerate(parsed):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=70, bottom=70, start=90, end=90)
            if i == 0:
                set_cell_shading(cell, LIGHT_FILL)
            text = row_data[j] if j < len(row_data) else ""
            p = cell.paragraphs[0]
            add_text_runs(p, text, default_bold=(i == 0), default_color=ACCENT if i == 0 else DARK)
            for run in p.runs:
                run.font.size = Pt(9.5 if cols >= 3 else 10)


def style_quote_paragraph(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE_FILL)
    ppr.append(shd)
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "4B8991")
    pbdr.append(left)
    ppr.append(pbdr)
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.right_indent = Cm(0.2)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(4)


def add_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table_from_lines(doc, table_lines)
            continue

        hm = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if hm:
            level = min(4, len(hm.group(1)))
            p = doc.add_heading(level=level)
            add_text_runs(p, hm.group(2), default_bold=True,
                          default_color=ACCENT if level <= 2 else DARK)
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph()
            style_quote_paragraph(p)
            content = stripped[1:].strip()
            add_text_runs(p, content, default_color=DARK)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.65)
            p.paragraph_format.first_line_indent = Cm(-0.25)
            add_text_runs(p, stripped[2:])
            i += 1
            continue

        nm = NUMBERED_RE.match(stripped)
        if nm and not stripped.startswith(tuple(f"{n}." for n in range(10, 100))):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.65)
            p.paragraph_format.first_line_indent = Cm(-0.25)
            add_text_runs(p, nm.group(2))
            i += 1
            continue

        p = doc.add_paragraph()
        add_text_runs(p, stripped)
        i += 1


def compose_markdown(module_path: Path) -> str:
    common = "\n\n".join(p.read_text(encoding="utf-8").strip() for p in COMMON_FILES)
    module = module_path.read_text(encoding="utf-8").strip()
    if "<!-- VERSION_MODULE -->" not in common:
        raise RuntimeError("VERSION_MODULE marker missing")
    common = common.replace("<!-- VERSION_MODULE -->", module)
    final = FINAL_FILE.read_text(encoding="utf-8").strip()
    return common + "\n\n" + final + "\n"


def add_document_end(doc: Document, meta: dict) -> None:
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=160, bottom=160, start=180, end=180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_runs(p, f"—— {meta['subtitle']}｜最终可复制Prompt结束 ——", default_bold=True, default_color=ACCENT)


def validate_docx(path: Path, meta: dict, markdown: str) -> List[str]:
    errors: List[str] = []
    if not path.exists() or path.stat().st_size < 35000:
        errors.append(f"file missing or too small: {path}")
        return errors
    try:
        with zipfile.ZipFile(path, "r") as zf:
            bad = zf.testzip()
            if bad:
                errors.append(f"zip CRC error: {bad}")
            xml = zf.read("word/document.xml").decode("utf-8")
            if "TOC \\o" in xml or "TOC \\h" in xml:
                errors.append("naked TOC field code found")
            if "\ufffd" in xml or "□" in xml:
                errors.append("replacement or missing glyph marker found")
    except Exception as exc:
        errors.append(f"zip validation failed: {exc}")
        return errors

    try:
        doc = Document(path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        required = [
            "内容数据库重建",
            "答案边界复核",
            "唯一答案四重验证",
            "阅读理解专项规则",
            "反幻觉与可追溯机制",
            "中文／日文字体双向审计",
            "静态目录",
            "共同终审清单",
            "最终执行命令",
        ]
        for phrase in required:
            if phrase not in full_text:
                errors.append(f"required phrase missing: {phrase}")
        if meta["key"] == "ruby":
            for phrase in ["含单字ルビ标音版专属规则", "答案速查汇总Ruby＝0", "お金"]:
                if phrase not in full_text:
                    errors.append(f"ruby requirement missing: {phrase}")
        else:
            for phrase in ["不含标音版专属规则", "Word原生Ruby节点：0", "全文禁止新增Ruby"]:
                if phrase not in full_text:
                    errors.append(f"no-ruby requirement missing: {phrase}")
        if len(doc.paragraphs) < 180:
            errors.append(f"unexpectedly low paragraph count: {len(doc.paragraphs)}")
    except Exception as exc:
        errors.append(f"python-docx reopen failed: {exc}")
    return errors


def build_one(meta: dict) -> Tuple[Path, List[str]]:
    markdown = compose_markdown(meta["module"])
    doc = Document()
    configure_document(doc, meta)
    add_cover(doc, meta)
    add_usage_and_navigation(doc, meta, markdown)
    add_markdown(doc, markdown)
    add_document_end(doc, meta)
    add_header_footer(doc, meta)
    out = OUTPUT / meta["filename"]
    doc.save(out)
    return out, validate_docx(out, meta, markdown)


def main() -> int:
    all_errors: List[str] = []
    report_lines = ["高考日语新题型配套解析最终Prompt｜DOCX生成与结构检查报告", ""]
    for meta in VERSIONS:
        out, errors = build_one(meta)
        report_lines.append(f"文件：{out.name}")
        report_lines.append(f"大小：{out.stat().st_size} bytes")
        if errors:
            report_lines.append("状态：FAIL")
            report_lines.extend(f"- {e}" for e in errors)
            all_errors.extend(f"{out.name}: {e}" for e in errors)
        else:
            report_lines.append("状态：PASS")
        report_lines.append("")

    report = OUTPUT / "validation_report.txt"
    report.write_text("\n".join(report_lines), encoding="utf-8")
    print(report.read_text(encoding="utf-8"))
    if all_errors:
        print("Validation failed", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
