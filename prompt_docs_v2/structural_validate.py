from __future__ import annotations

import sys
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "output"

REQUIRED_COMMON = [
    "内容数据库重建",
    "答案边界复核",
    "唯一答案四重验证",
    "阅读理解专项规则",
    "反幻觉与可追溯机制",
    "中文／日文字体双向审计",
    "静态目录",
    "共同终审清单",
    "最终执行命令",
]


def validate(path: Path) -> list[str]:
    errors: list[str] = []
    if not path.exists() or path.stat().st_size < 35000:
        return ["文件不存在或体积异常"]
    try:
        with zipfile.ZipFile(path, "r") as zf:
            bad = zf.testzip()
            if bad:
                errors.append(f"ZIP CRC错误：{bad}")
            xml = zf.read("word/document.xml").decode("utf-8")
            if "TOC \\o" in xml or "TOC \\h" in xml:
                errors.append("发现裸露TOC域代码")
            if "\ufffd" in xml:
                errors.append("发现Unicode替代字符")
    except Exception as exc:
        return [f"DOCX压缩结构读取失败：{exc}"]

    try:
        doc = Document(path)
        text_parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.extend(p.text for p in cell.paragraphs)
        full_text = "\n".join(text_parts)
        for phrase in REQUIRED_COMMON:
            if phrase not in full_text:
                errors.append(f"缺少关键内容：{phrase}")
        if "含单字ルビ标音版" in path.name:
            for phrase in ["含单字ルビ标音版专属规则", "答案速查汇总Ruby＝0", "お金"]:
                if phrase not in full_text:
                    errors.append(f"含标音版缺少：{phrase}")
        if "不含标音版" in path.name:
            for phrase in ["不含标音版专属规则", "Word原生Ruby节点：0", "全文禁止新增Ruby"]:
                if phrase not in full_text:
                    errors.append(f"不含标音版缺少：{phrase}")
        if len(doc.paragraphs) < 180:
            errors.append(f"段落数量异常：{len(doc.paragraphs)}")
    except Exception as exc:
        errors.append(f"python-docx重新打开失败：{exc}")
    return errors


def main() -> int:
    docs = sorted(OUTPUT.glob("*.docx"))
    lines = ["高考日语新题型配套解析最终Prompt｜DOCX结构检查报告", ""]
    all_errors: list[str] = []
    if len(docs) != 2:
        all_errors.append(f"应有2份DOCX，实际{len(docs)}份")
    for path in docs:
        errors = validate(path)
        lines.append(f"文件：{path.name}")
        lines.append(f"大小：{path.stat().st_size if path.exists() else 0} bytes")
        if errors:
            lines.append("状态：FAIL")
            lines.extend(f"- {e}" for e in errors)
            all_errors.extend(f"{path.name}: {e}" for e in errors)
        else:
            lines.append("状态：PASS")
        lines.append("")
    report = OUTPUT / "validation_report.txt"
    report.write_text("\n".join(lines), encoding="utf-8")
    print(report.read_text(encoding="utf-8"))
    if all_errors:
        print("\n".join(all_errors), file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
